# --- START OF FILE file_processor.py ---
import fitz
import re
import nltk
import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH # Import alignment enum
import streamlit as st

# --- NLTK Download Logic ---
# (Keep as is)
def download_nltk_data(resource_name, resource_path):
    try: nltk.data.find(resource_path)
    except LookupError:
        st.info(f"Downloading NLTK data package: '{resource_name}'...")
        try: nltk.download(resource_name, quiet=True); st.success(f"NLTK data '{resource_name}' downloaded."); return True
        except Exception as e: st.error(f"Download Error: Failed for NLTK '{resource_name}'. Error: {e}"); return False
    except Exception as e_find: st.error(f"NLTK Find Error ({resource_name}): {e_find}"); return False
    return True

# --- Metadata/Footer Check ---
# (Keep as is)
def is_likely_metadata_or_footer(line):
    line = line.strip();
    if not line: return True
    cleaned_line = re.sub(r"^\W+|\W+$", "", line)
    if cleaned_line.isdigit() and line == cleaned_line and len(cleaned_line) < 4 : return True
    if "www." in line or ".com" in line or "@" in line or "books" in line.lower() or "global" in line.lower() or ("center" in line.lower() and "peace" in line.lower()):
        if len(line.split()) < 12: return True
    if re.match(r"^\s*Page\s+\d+\s*$", line, re.IGNORECASE): return True
    if "©" in line or "copyright" in line.lower() or "first published" in line.lower() or "isbn" in line.lower(): return True
    if "printed in" in line.lower(): return True
    if any(loc in line.lower() for loc in ["nizamuddin", "new delhi", "noida", "bensalem", "byberry road"]): return True
    if len(set(line.strip())) < 4 and len(line.strip()) > 4: return True
    if line.endswith(cleaned_line) and cleaned_line.isdigit() and len(line) < 80 and len(line) > len(cleaned_line) + 2:
         if not re.search(r"[a-zA-Z]{4,}", line): return True
    if len(line) < 15 and line.isupper() and not re.match(r"^\s*(CHAPTER|SECTION|PART)\s", line, re.IGNORECASE):
        if len(line.split()) > 1 and not line.istitle() : return True
    return False

# --- HEADING CHECKER v16 ---
def check_heading_user_defined(
    line_dict, line_text, page_width, is_single_line_block,
    is_bold_hint, is_italic_hint, is_centered_hint, # Pass centering hint now
    heading_criteria ):
    """ Checks if a line matches user-defined heading criteria. """
    words = line_text.split(); num_words = len(words)
    if not line_text or num_words == 0: return None
    if is_likely_metadata_or_footer(line_text): return None

    # --- Apply User Selected Criteria ---
    # Length
    if heading_criteria['use_length']:
        if not (heading_criteria['min_words'] <= num_words <= heading_criteria['max_words']): return None
    # Keyword
    if heading_criteria['keyword_pattern']:
        try:
            if not re.search(heading_criteria['keyword_pattern'], line_text, re.IGNORECASE): return None
        except re.error: return None
    # Isolation (PDF only)
    if heading_criteria['require_isolated'] and line_dict is not None and not is_single_line_block: return None
    # Centering (Use hint provided by caller)
    if heading_criteria['require_centered'] and not is_centered_hint: return None

    # --- Style & Case Checks ---
    line_is_italic = is_italic_hint; line_is_bold = is_bold_hint
    if line_dict: # If PDF, refine style checks using spans
        try:
            valid_spans = [s for s in line_dict["spans"] if s['text'].strip()]
            if valid_spans:
                total_chars=sum(len(s['text'].strip()) for s in valid_spans); italic_chars=0; bold_chars=0
                for s in valid_spans:
                    flags=s.get('flags',0); font_name=s.get('font','').lower(); span_len=len(s['text'].strip())
                    if flags & 1 or "italic" in font_name: italic_chars += span_len
                    if flags & 4 or "bold" in font_name or "black" in font_name: bold_chars += span_len
                if total_chars > 0: line_is_italic=(italic_chars/total_chars)>0.6; line_is_bold=(bold_chars/total_chars)>0.6
        except Exception: pass
    if heading_criteria['require_italic'] and not line_is_italic: return None
    if heading_criteria['require_bold'] and not line_is_bold: return None

    is_line_title_case = line_text.istitle(); is_line_all_caps = line_text.isupper() and re.search("[A-Z]", line_text)
    if heading_criteria['require_title_case'] and not is_line_title_case: return None
    if heading_criteria['require_all_caps'] and not is_line_all_caps: return None
    if heading_criteria['require_title_case'] and heading_criteria['require_all_caps']:
        if not (is_line_title_case and is_line_all_caps):
             if num_words > 1: return None

    # If all required checks passed:
    # print(f"✅ Matched Heading: {line_text}") # Debug
    return line_text # Return the heading text


# --- Main Extraction Function ---
def extract_sentences_with_structure(
    file_name, file_content, heading_criteria,
    start_skip=0, end_skip=0, start_page_offset=1 ):
    """ Extracts text, handles PDF/DOCX, detects headings based on criteria. """
    _ALIGN_PARAGRAPH # Import alignment constants
import streamlit as st

# --- NLTK Download Logic ---
# (Keep as is)
def download_nltk_data(resource_name, resource_path):
    try: nltk.data.find(resource_path)
    except LookupError:
        st.info(f"Downloading NLTK data package: '{resource_name}'...")
        try: nltk.download(resource_name, quiet=True); st.success(f"NLTK data '{resource_name}' downloaded."); return True
        except Exception as e: st.error(f"Download Error: Failed for NLTK '{resource_name}'. Error: {e}"); return False
    except Exception as e_find: st.error(f"NLTK Find Error ({resource_name}): {e_find}"); return False
    return True

# --- Metadata/Footer Check ---
# (Keep as is)
def is_likely_metadata_or_footer(line):
    line = line.strip();
    if not line: return True
    cleaned_line = re.sub(r"^\W+|\W+$", "", line)
    if cleaned_line.isdigit() and line == cleaned_line and len(cleaned_line) < 4 : return True
    if "www." in line or ".com" in line or "@" in line or "books" in line.lower() or "global" in line.lower() or ("center" in line.lower() and "peace" in line.lower()):
        if len(line.split()) < 12: return True
    if re.match(r"^\s*Page\s+\d+\s*$", line, re.IGNORECASE): return True
    if "©" in line or "copyright" in line.lower() or "first published" in line.lower() or "isbn" in line.lower(): return True
    if "printed in" in line.lower(): return True
    if any(loc in line.lower() for loc in ["nizamuddin", "new delhi", "noida", "bensalem", "byberry road"]): return True
    if len(set(line.strip())) < 4 and len(line.strip()) > 4: return True
    if line.endswith(cleaned_line) and cleaned_line.isdigit() and len(line) < 80 and len(line) > len(cleaned_line) + 2:
         if not re.search(r"[a-zA-Z]{4,}", line): return True
    if len(line) < 15 and line.isupper() and not re.match(r"^\s*(CHAPTER|SECTION|PART)\s", line, re.IGNORECASE):
        if len(line.split()) > 1 and not line.istitle() : return True
    return False


# --- Heading Checker (Handles PDF dict or DOCX para info) ---
def check_heading_user_defined(
    line_dict,              # Dictionary for PDF line (contains bbox, spans) OR None for DOCX
    line_text,              # Plain text of the line/paragraph
    page_width,             # Width of the PDF page (0 for DOCX)
    is_single_line_block,   # Is this line alone in its block (PDF)? False for DOCX
    is_bold_hint,           # Is the DOCX paragraph bold?
    is_italic_hint,         # Is the DOCX paragraph italic?
    para_alignment,         # DOCX Paragraph alignment enum (or None)
    heading_criteria        # Dictionary of user selections from UI
    ):
    """ Checks if a line/paragraph matches user-defined heading criteria. """
    words = line_text.split()
    num_words = len(words)
    if not line_text or num_words == 0: return None
    if is_likely_metadata_or_footer(line_text): return None

    # --- Apply User Selected Criteria ---
    if heading_criteria['use_length']:
        if not (heading_criteria['min_words'] <= num_words <= heading_criteria['max_words']): return None
    if heading_criteria['keyword_pattern']:
        try:
            if not re.search(heading_criteria['keyword_pattern'], line_text, re.IGNORECASE): return None
        except re.error: return None

    # --- Apply PDF Specific Checks (if line_dict is provided) ---
    if line_dict:
        if heading_criteria['require_isolated'] and not is_single_line_block: return None
        if heading_criteria['require_centered']:
            is_centered = False
            bbox = line_dict.get('bbox', None)
            if bbox and page_width > 0:
                left = bbox[0]; right = bbox[2]
                if abs(left - (page_width - right)) < (page_width * 0.20) and left > (page_width * 0.15): is_centered = True
            if not is_centered: return None # Fail if required but not met

        # Extract PDF style info
        line_is_italic_pdf = False; line_is_bold_pdf = False
        try:
            valid_spans = [s for s in line_dict["spans"] if s['text'].strip()]
            if valid_spans:
                total_chars = sum(len(s['text'].strip()) for s in valid_spans)
                italic_chars = 0; bold_chars = 0
                for s in valid_spans:
                    flags = s.get('flags', 0); font_name = s.get('font','').lower(); span_len = len(s['text'].strip())
                    if flags & 1 or "italic" in font_name: italic_chars += span_len
                    if flags & 4 or "bold" in font_name or "black" in font_name: bold_chars += span_len
                if total_chars > 0: line_is_italic_extracted_data = []; current_chapter_title_state = None; doc = None
    file_extension = file_name.split('.')[-1].lower()

    if file_extension == 'pdf':
        try:
            doc = fitz.open(stream=file_content, filetype="pdf"); total_pages = len(doc)
            for page_num_0based, page in enumerate(doc):
                if page_num_0based < start_skip: continue
                if page_num_0based >= total_pages - end_skip: break
                adjusted_page_num = page_num_0based - start_skip + start_page_offset
                page_marker = adjusted_page_num; page_width = page.rect.width
                try:
                    blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT | fitz.TEXT_PRESERVE_LIGATURES)["blocks"]
                    for b in blocks:
                        if b['type'] == 0:
                            is_single_line_block = len(b['lines']) == 1
                            for l in b["lines"]:
                                line_dict = l; line_text = "".join(s["text"] for s in l["spans"]).strip()
                                if not line_text or is_likely_metadata_or_footer(line_text): continue

                                # Centering hint for PDF
                                is_centered_hint = False
                                bbox = l.get('bbox');
                                if bbox and page_width > 0:
                                    left=bbox[0]; right=bbox[2]
                                    if abs(left - (page_width - right)) < (page_width * 0.20) and left > (page_width * 0.15):
                                        is_centered_hint = True

                                heading_text = check_heading_user_defined(
                                    line_dict, line_text, page_width, is_single_line_block,
                                    False, False, # is_bold/is_italic hints are extracted inside checker for PDF
                                    is_centered_hint, # Pass centering hint
                                    heading_criteria
                                )
                                if heading_text is not None:
                                    current_chapter_title_state = heading_text
                                    extracted_data.append((heading_text, page_marker, heading_text)) # (text, marker, chapter_title)
                                else:
                                    try:
                                        sentences = nltk.sent_tokenize(line_text)
                                        for sentence in sentences:
                                            sc = sentence.strip();
                                            if sc: extracted_data.append((sc, page_marker, None))
                                    except Exception as e: print(f"Warn: NLTK err PDF Pg {page_marker}: {e}"); extracted_data.append((line_text, page_marker, None))
                except Exception as e_page: print(f"Error processing PDF page {adjusted_page_num}: {e_page}")
        except Exception as e_main: print(f"Main PDF Extraction Error: {e_main}"); return None
        finally:
            if doc: doc.close()

    elif file_extension == 'docx':
        try:
            document = docx.Document(io.BytesIO(file_content)); paragraph_index = 0
            for para in document.paragraphs:
                paragraph_index += 1; page_marker = f"Para_{paragraph_index}"
                line_text = para.text.strip()
                if not line_text or is_likely_metadata_or_footer(line_text): continue

                # Get hints from DOCX paragraph properties
                is_bold_hint = any(run.bold for run in para.runs if run.text.strip())
                is_italic_hint = any(run.italic for run in para.runs if run.text.strip())
                # Check alignment for centering hint
                is_centered_hint = (para.alignment == WD_ALIGN_PARAGRAPH.CENTER) if para.alignment else False

                heading_text = check_heading_user_defined(
                    None, line_text, 0, False, # No PDF-specific info for DOCX
                    is_bold_hint, is_italic_hint, is_centered_hint, # Pass hints
                    heading_criteria
                )
                if heading_text is not None:
                    current_chapter_title_state = heading_text
                    extracted_data.append((heading_text, page_marker, heading_text))
                else:
                    try:
                        sentences = nltk.sent_tokenize(line_text)
                        for sentence in sentences:
                            sc = sentence.strip();
                            if sc: extracted_data.append((sc, page_marker, None))
                    except Exception as e: print(f"Warn: NLTK err DOCX Para {page_marker}: {e}"); extracted_data.append((line_text, page_marker,pdf = (italic_chars / total_chars) > 0.6; line_is_bold_pdf = (bold_chars / total_chars) > 0.6
        except Exception: pass
        if heading_criteria['require_italic'] and not line_is_italic_pdf: return None
        if heading_criteria['require_bold'] and not line_is_bold_pdf: return None

    # --- Apply DOCX Specific Checks (if line_dict is None) ---
    else: # DOCX mode
        if heading_criteria['require_isolated']: pass # Cannot check reliably for DOCX paragraphs
        if heading_criteria['require_centered']:
            # Check paragraph alignment enum passed from DOCX processing
            if para_alignment != WD_ALIGN_PARAGRAPH.CENTER:
                return None # Fails centering constraint for DOCX
        if heading_criteria['require_italic'] and not is_italic_hint: return None # Use hints passed for DOCX
        if heading_criteria['require_bold'] and not is_bold_hint: return None

    # --- Apply Case Checks (Common to both) ---
    is_line_title_case = line_text.istitle()
    is_line_all_caps = line_text.isupper() and re.search("[A-Z]", line_text)
    if heading_criteria['require_title_case'] and not is_line_title_case: return None
    if heading_criteria['require_all_caps'] and not is_line_all_caps: return None
    if heading_criteria['require_title_case'] and heading_criteria['require_all_caps']:
        if not (is_line_title_case and is_line_all_caps):
             if num_words > 1: return None

    # If all active required checks passed:
    # print(f"✅ Matched Heading: {line_text}") # Debug
    # Simplification: Return only 'chapter' type for now
    return line_text


# --- Main Extraction Function ---
def extract_sentences_with_structure(
    file_name, file_content,
    heading_criteria, # Pass the dictionary of user choices
    start_skip=0, end_skip=0, start_page_offset=1
    ):
    extracted_data = []
    current_chapter_title_state = None
    doc = None

    file_extension = file_name.split('.')[-1].lower()

    # --- PDF Processing ---
    if file_extension == 'pdf':
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total_pages = len(doc)
            for page_num_0based, page in enumerate(doc):
                # ...(page skipping)...
                if page_num_0based < start_skip: continue
                if page_num_0based >= total_pages - end_skip: break
                adjusted_page_num = page_num_0based - start_skip + start_page_offset
                page_marker = adjusted_page_num
                page_width = page.rect.width
                try:
                    blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT | fitz.TEXT_PRESERVE_LIGATURES)["blocks"]
                    for b in blocks:
                        if b['type'] == 0:
                            is_single_line_block = len(b['lines']) == 1
                            for l in b["lines"]:
                                line_dict = l
                                line_text_raw = "".join(s["text"] for s in l["spans"]).strip()
                                if not line_text_raw or is_likely_metadata_or_footer(line_text_raw): continue

                                # Check heading
                                heading_text = check_heading_user_defined(
                                    line_dict, line_text_raw, page_width, is_single_line_block,
                                    False, False, None, # Pass False/None for DOCX-specific params
                                    heading_criteria
                                )

                                if heading_text is not None:
                                    current_chapter_title_state = heading_text
                                    extracted_data.append((heading_text, page_marker, heading_text))
                                else: # Regular text
                                    try: # Tokenize
                                        sentences = nltk.sent_tokenize(line_text_raw)
                                        for sentence in sentences:
                                            sc = sentence.strip();
                                            if sc: extracted_data.append((sc, page_marker, None))
                                    except Exception as e_nltk: print(f"Warn: NLTK PDF {page_marker}: {e_nltk}"); extracted_data.append((line_text_raw, page_marker, None))
                except Exception as e_page: print(f"Error processing PDF page {adjusted_page_num}: {e_page}")
        except Exception as e_main: print(f"Main PDF Error: {e_main}"); return None
        finally:
            if doc: doc.close()

    # --- DOCX Processing ---
    elif file_extension == 'docx':
        try:
            document = docx.Document(io.BytesIO(file_content))
            paragraph_index = 0
            for para in document.paragraphs:
                paragraph_index += 1
                page_marker = f"Para_{paragraph_index}"
                line_text = para.text.strip()
                if not line_text or is_likely_metadata_or_footer(line_text): continue

                is_bold_hint = any(run.bold for run in para.runs if run.text.strip())
                is_italic_hint = any(run.italic for run in para.runs if run.text.strip())
                # Get alignment (default to LEFT if not set)
                para_alignment = para.alignment if para.alignment is not None else WD_ALIGN_PARAGRAPH.LEFT

                # Check heading
                heading_text = check_heading_user_defined(
                    None, line_text, 0, False, # Pass None/False for PDF-specific params
                    is_bold_hint, is_italic_hint, # Pass style hints
                    para_alignment, # Pass alignment
                    heading_criteria
                )

                if heading_text is not None:
                    current_chapter_title_state = heading_text
                    extracted_data.append((heading_text, page_marker, heading_text))
                else: # Regular text
                    try: # Tokenize
                        sentences = nltk.sent_tokenize(line_text)
                        for sentence in sentences:
                            sc = sentence.strip()
                            if sc: extracted_data.append((sc, page_marker, None))
                    except Exception as e_nltk: print(f"Warn: NLTK DOCX {page_marker}: {e_nltk}"); extracted_data.append((line_text, page_marker, None))
        except Exception as e_main: print(f"Main DOCX Error: {e_main}"); return None

    # --- Unsupported ---
    else: print(f"Error: Unsupported file type: .{file_extension}"); return None

    print(f"Extraction complete. Found {len(extracted_data)} items.")
    return extracted_data

# --- END OF FILE file_processor.py ---
